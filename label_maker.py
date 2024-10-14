import os
import platform
import customtkinter as ctk
from tkinter import filedialog, messagebox, colorchooser
from PIL import Image
from docxtpl import DocxTemplate, RichText
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sys

# Function to locate resource files, works for both PyInstaller executable and dev environment
def resource_path(relative_path):
    """ Get the absolute path to a resource, works for dev and for PyInstaller """
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# Set CustomTkinter appearance mode and color theme
ctk.set_appearance_mode("light")  # Modes: "System" (default), "Dark", "Light"
ctk.set_default_color_theme("green")  # Themes: "blue", "dark-blue", "green"

# Global variable to store label data
label_data = []

# Paths for docx files
blank_template_path = resource_path('Label_Template_BLANK.docx')
generated_template_path = resource_path('GENERATED_Label_Template.docx')

# Function to create labels in the DOCX file with font size 16 and bold formatting (excluding curly braces)
def create_labels(doc, start_number, num_labels):
    for i in range(start_number, start_number + num_labels):
        order_paragraph = add_label_paragraph(doc, f"Order Name & Number\n{{{{ order_name{i} }}}}")
        batch_chip_paragraph = add_label_paragraph(doc, f"Batch / Chip Number\n{{{{ batch_chip{i} }}}}")
        type_paragraph = add_label_paragraph(doc, f"Type: {{{{ card_envelope{i} }}}}")

# Function to add paragraphs and apply bold formatting with 16-point font size
def add_label_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    is_bold = True  # Start by setting bold for non-curly brace text
    current_run = ""
    
    for idx, char in enumerate(text):
        if char == '{' and text[idx+1] == '{':  # Detect start of curly brace
            if current_run:
                run = paragraph.add_run(current_run)
                run.bold = is_bold
                run.font.size = Pt(16)  # Set to 16-point font size
            current_run = char
            is_bold = False  # Inside curly braces, not bold
        elif char == '}' and text[idx-1] == '}':  # Detect end of curly brace
            current_run += char
            run = paragraph.add_run(current_run)
            run.bold = is_bold
            run.font.size = Pt(16)  # Set to 16-point font size
            current_run = ""
            is_bold = True  # After closing curly braces, return to bold
        else:
            current_run += char

    # Add any remaining text after the loop
    if current_run:
        run = paragraph.add_run(current_run)
        run.bold = is_bold
        run.font.size = Pt(16)  # Set to 16-point font size
    
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    return paragraph

# Function to create and save the DOCX file with stored label information
def create_docx():
    if not label_data:
        messagebox.showerror("Error", "No label data to create the file!")
        return

    try:
        # Load the blank label template and delete the first line
        doc = Document(blank_template_path)
        if len(doc.paragraphs) > 0:
            doc.paragraphs[0]._element.getparent().remove(doc.paragraphs[0]._element)

        # Determine how many labels we need
        num_labels_to_generate = len(label_data)
        create_labels(doc, 1, num_labels_to_generate)

        # Save the updated document template with placeholders
        doc.save(generated_template_path)

        # Now use this generated template in docxtpl to replace placeholders with real data
        template = DocxTemplate(generated_template_path)
        context = {}
        for i, label in enumerate(label_data):
            context[f'order_name{i+1}'] = RichText(label['order_name'], color=label['color'], size=32, bold=True)
            context[f'batch_chip{i+1}'] = RichText(label['batch_chip'], color=label['color'], size=32, bold=True)
            context[f'card_envelope{i+1}'] = RichText(label['card_envelope'], color=label['color'], size=32, bold=True)


        # Ask user to save the file
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        
        if save_path:
            template.render(context)
            template.save(save_path)
            messagebox.showinfo("Success", f"Labels saved to {save_path}")
            # Show the "Open File" button
            open_button.pack(pady=10, padx=20)
            open_button.configure(command=lambda: open_docx_file(save_path))
        else:
            messagebox.showerror("Error", "Save path not specified or operation cancelled.")
    
    except Exception as e:
        messagebox.showerror("Error", f"Failed to create the DOCX file: {str(e)}")

# Function to add label information to the label_data list and display it
def add_label_data():
    order_name = order_name_entry.get()
    num_hw_machines = num_hw_machines_entry.get()
    card_envelope = card_envelope_var.get()

    if not order_name or not num_hw_machines or not card_envelope:
        messagebox.showerror("Error", "All fields must be filled in!")
        return

    try:
        num_hw_machines = int(num_hw_machines)
    except ValueError:
        messagebox.showerror("Error", "Number of HW Machines must be a valid number!")
        return

    # Check for duplicates based on both order_name and card_envelope (Envelopes or Cards)
    if any(label['order_name'] == order_name and label['card_envelope'] == card_envelope for label in label_data):
        messagebox.showerror("Error", f"Label for '{order_name} {card_envelope}' already exists!")
        return

    # Ask user to choose a color for the label
    color = colorchooser.askcolor(title="Choose Label Color")[1]  # Get the hex color value
    if not color:
        messagebox.showerror("Error", "No color selected!")
        return

    # Prepare the label data
    for i in range(1, num_hw_machines + 1):
        label_data.append({
            'order_name': order_name,
            'batch_chip': f"{i} of {num_hw_machines}",
            'card_envelope': card_envelope,
            'color': color[1:]  # Remove the "#" symbol from the hex color
        })
    
    # Display the added label with color and make it clickable to change the color
    display_label(f"{order_name} {card_envelope}", color)

# Function to display the label data and allow color changes
def display_label(order_name, color):
    # Create the label with the assigned color
    color_label = ctk.CTkLabel(scrollable_frame, text=f"{order_name} assigned color", fg_color=color, text_color="white", width=300)
    color_label.pack(pady=5, padx=20, anchor="w")

    # Bind a click event to the label to change the color
    color_label.bind("<Button-1>", lambda e: change_color(order_name, color_label))

# Function to change the color when the label is clicked
def change_color(order_name, color_label):
    # Open the color chooser dialog
    new_color = colorchooser.askcolor(title=f"Choose a new color for {order_name}")[1]
    
    if new_color:
        # Update the color label's background color
        color_label.configure(fg_color=new_color)
        
        # Update the label_data for this order_name with the new color
        for label in label_data:
            if f"{label['order_name']} {label['card_envelope']}" == order_name:
                label['color'] = new_color[1:]  # Strip the "#" from the hex color

# Function to open the DOCX file
def open_docx_file(file_path):
    try:
        if platform.system() == "Windows":
            os.startfile(file_path)
        elif platform.system() == "Darwin":
            os.system(f"open {file_path}")
        else:
            os.system(f"xdg-open {file_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Failed to open the DOCX file: {str(e)}")

# Function to reset label data and clear the displayed labels
def reset_labels():
    global label_data

    # Clear the label data list
    label_data.clear()

    # Destroy all the labels inside the scrollable frame
    for widget in scrollable_frame.winfo_children():
        widget.destroy()

    # Show a message box confirming reset
    messagebox.showinfo("Reset", "All label data has been reset!")

# Initialize the app
root = ctk.CTk()
root.title("Manual Label Maker")
root.geometry("700x700")
root.configure(bg="#3A3A3A")

# Load the custom icon in .ico format (replace the default CustomTkinter icon)
icon_path = resource_path('scribe-icon.ico')  # Ensure the icon is in .ico format
root.iconbitmap(icon_path)  # Set the custom icon for the window


# Load the logo image using PIL
logo_path = resource_path('scribe-logo-final.png')  # Ensure the logo is in the same folder or packaged
logo_image = Image.open(logo_path)  # Load the image
logo_image = logo_image.resize((258, 100), Image.Resampling.LANCZOS)  # Resize the image

# Convert the PIL image to a CTkImage to avoid the HighDPI warning
logo_ctk_image = ctk.CTkImage(light_image=logo_image, dark_image=logo_image, size=(258, 100))

# Add the logo to a label at the top of the window
logo_label = ctk.CTkLabel(root, image=logo_ctk_image, text="")  # Display only the image, no text
logo_label.pack(pady=10)  # Adjust padding as needed

# Create a scrollable frame for displaying added labels
canvas = ctk.CTkCanvas(root, bg="#3A3A3A")
scrollbar = ctk.CTkScrollbar(root, orientation="vertical", command=canvas.yview)
scrollable_frame = ctk.CTkFrame(canvas, fg_color="#3A3A3A", corner_radius=15)

scrollable_frame.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))

canvas.create_window((0, 0), window=scrollable_frame, anchor="nw", width=680)
canvas.configure(yscrollcommand=scrollbar.set)

canvas.pack(side="left", fill="both", expand=True)
scrollbar.pack(side="right", fill="y")

# Input fields for manual label entry
order_name_label = ctk.CTkLabel(root, text="Order Name:", font=("Helvetica", 12), text_color="black")
order_name_label.pack(pady=5, padx=20, anchor="w")
order_name_entry = ctk.CTkEntry(root, width=300)
order_name_entry.pack(pady=5, padx=20)

num_hw_machines_label = ctk.CTkLabel(root, text="Number of HW Machines:", font=("Helvetica", 12), text_color="black")
num_hw_machines_label.pack(pady=5, padx=20, anchor="w")
num_hw_machines_entry = ctk.CTkEntry(root, width=300)
num_hw_machines_entry.pack(pady=5, padx=20)

# Variable to store selected card/envelope type
card_envelope_var = ctk.StringVar(value="Envelopes")

# Radio buttons for "Envelopes" and "Cards"
radio_envelopes = ctk.CTkRadioButton(root, text="Envelopes", variable=card_envelope_var, value="Envelopes")
radio_cards = ctk.CTkRadioButton(root, text="Cards", variable=card_envelope_var, value="Cards")

# Packing the radio buttons below inputs
radio_envelopes.pack(pady=10, padx=20, anchor="w")
radio_cards.pack(pady=10, padx=20, anchor="w")

# Button to add label data
add_button = ctk.CTkButton(root, text="Add to Label File", command=add_label_data, width=300, fg_color="#133d8e", hover_color="#266cc3")
add_button.pack(pady=20, padx=20)

# GUI Setup for the "Reset" button
reset_button = ctk.CTkButton(root, text="Reset", command=reset_labels, width=300, fg_color="#8e1313", hover_color="#c32626")
reset_button.pack(pady=20, padx=20)

# Button to create DOCX
create_button = ctk.CTkButton(root, text="Create DOCX", command=create_docx, width=300, fg_color="#133d8e", hover_color="#266cc3")
create_button.pack(pady=20, padx=20)

# Add a hidden placeholder for the "Open File" button
open_button = ctk.CTkButton(root, text="Open Created DOCX File", width=300, fg_color="#133d8e", hover_color="#266cc3")
open_button.pack_forget()  # Initially hide the button

# Main GUI loop
root.mainloop()
